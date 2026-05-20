"""文档加载与解析模块。

支持 PDF、TXT、Markdown 格式，统一输出结构化文档内容 + metadata。
包含文本清洗逻辑。
"""

import re
import unicodedata
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

try:
    import pymupdf as fitz
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None


@dataclass
class Document:
    """统一文档数据结构。"""

    doc_id: str
    title: str
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: list[dict[str, Any]] = field(default_factory=list)

    def __post_init__(self) -> None:
        if not self.metadata.get("created_at"):
            self.metadata["created_at"] = datetime.now().isoformat()


class DocumentParser(ABC):
    """文档解析器抽象基类。"""

    @abstractmethod
    def parse(self, file_path: str) -> Document:
        """解析文档并返回 Document 对象。"""
        pass


class PDFParser(DocumentParser):
    """PDF 文档解析器。使用 PyMuPDF (fitz)。"""

    def __init__(self, extract_images: bool = False):
        self.extract_images = extract_images

    def parse(self, file_path: str) -> Document:
        if fitz is None:
            raise ImportError("请安装 pymupdf: pip install pymupdf")

        doc_id = self._generate_doc_id(file_path)
        sections: list[dict[str, Any]] = []

        with fitz.open(file_path) as pdf:
            title = self._extract_title(pdf) or Path(file_path).stem
            full_text_parts: list[str] = []

            for page_num, page in enumerate(pdf):
                text = page.get_text("text") or ""
                if text.strip():
                    sections.append({
                        "page": page_num + 1,
                        "text": text.strip(),
                        "type": "page",
                    })
                    full_text_parts.append(text.strip())

            full_text = "\n\n".join(full_text_parts)

        return Document(
            doc_id=doc_id,
            title=title,
            content=full_text,
            metadata={
                "source": str(file_path),
                "file_type": "pdf",
                "total_pages": len(pdf) if fitz else 0,
                "created_at": self._extract_creation_date(pdf) if fitz else None,
            },
            sections=sections,
        )

    def _extract_title(self, pdf: Any) -> Optional[str]:
        """尝试从 PDF 元数据或第一行提取标题。"""
        metadata = pdf.metadata
        if metadata and metadata.get("title"):
            return metadata["title"]

        first_page_text = ""
        if len(pdf) > 0:
            first_page_text = pdf[0].get_text("text") or ""

        lines = [l.strip() for l in first_page_text.split("\n") if l.strip()]
        if lines:
            return lines[0][:200]

        return None

    def _extract_creation_date(self, pdf: Any) -> Optional[str]:
        metadata = pdf.metadata
        if metadata and metadata.get("creationDate"):
            return metadata["creationDate"]
        return None

    @staticmethod
    def _generate_doc_id(file_path: str) -> str:
        import hashlib
        path_bytes = file_path.encode("utf-8")
        return f"doc_{hashlib.md5(path_bytes).hexdigest()[:12]}"


class MarkdownParser(DocumentParser):
    """Markdown 文档解析器。"""

    def parse(self, file_path: str) -> Document:
        doc_id = PDFParser._generate_doc_id(file_path)
        with open(file_path, encoding="utf-8") as f:
            raw_text = f.read()

        title = self._extract_title_from_md(raw_text) or Path(file_path).stem
        sections = self._extract_sections(raw_text)

        cleaned_content = self._clean_text(raw_text)

        return Document(
            doc_id=doc_id,
            title=title,
            content=cleaned_content,
            metadata={
                "source": str(file_path),
                "file_type": "md",
            },
            sections=sections,
        )

    def _extract_title_from_md(self, content: str) -> Optional[str]:
        match = re.match(r"^#\s+(.+)$", content, re.MULTILINE)
        if match:
            return match.group(1).strip()
        return None

    def _extract_sections(self, content: str) -> list[dict[str, Any]]:
        """按 Markdown 标题层级提取章节。"""
        sections = []
        lines = content.split("\n")
        current_section = {"level": 0, "title": "", "content": [], "type": "section"}

        for line in lines:
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                if current_section["content"]:
                    sections.append({
                        "level": current_section["level"],
                        "title": current_section["title"],
                        "text": "\n".join(current_section["content"]).strip(),
                        "type": "section",
                    })
                level = len(header_match.group(1))
                current_section = {
                    "level": level,
                    "title": header_match.group(2).strip(),
                    "content": [],
                    "type": "section",
                }
            else:
                current_section["content"].append(line)

        if current_section["content"]:
            sections.append({
                "level": current_section["level"],
                "title": current_section["title"],
                "text": "\n".join(current_section["content"]).strip(),
                "type": "section",
            })

        return sections

    def _clean_text(self, text: str) -> str:
        lines = text.split("\n")
        cleaned_lines = [self._clean_line(line) for line in lines]
        return "\n".join(cleaned_lines).strip()

    @staticmethod
    def _clean_line(line: str) -> str:
        line = line.strip()
        line = unicodedata.normalize("NFKC", line)
        line = re.sub(r"[ \t]+", " ", line)
        return line


class TextParser(DocumentParser):
    """纯文本文件解析器。"""

    def parse(self, file_path: str) -> Document:
        doc_id = PDFParser._generate_doc_id(file_path)
        with open(file_path, encoding="utf-8") as f:
            raw_text = f.read()

        cleaned_content = self._clean_text(raw_text)

        return Document(
            doc_id=doc_id,
            title=Path(file_path).stem,
            content=cleaned_content,
            metadata={
                "source": str(file_path),
                "file_type": "txt",
            },
            sections=[{"type": "section", "title": "全文", "text": cleaned_content}],
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            line = unicodedata.normalize("NFKC", line)
            line = re.sub(r"[ \t]+", " ", line)
            if line:
                cleaned_lines.append(line)
        return "\n".join(cleaned_lines)


class DocumentLoader:
    """统一文档加载器，自动根据文件扩展名选择解析器。

    使用示例:
        loader = DocumentLoader()
        docs = loader.load_directory("./data/documents")
        for doc in docs:
            print(f"{doc.title}: {len(doc.content)} chars, {len(doc.sections)} sections")
    """

    _parsers: dict[str, type[DocumentParser]] = {
        ".pdf": PDFParser,
        ".md": MarkdownParser,
        ".txt": TextParser,
    }

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def load(self, file_path: str) -> Document:
        """加载单个文档。"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        parser_cls = self._parsers.get(suffix)
        if not parser_cls:
            raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: {list(self._parsers.keys())}")

        parser = parser_cls()
        return parser.parse(file_path)

    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        tenant_id: Optional[str] = None,
        department: Optional[str] = None,
        author: Optional[str] = None,
        doc_type: Optional[str] = None,
    ) -> list[Document]:
        """批量加载目录下所有支持的文档。

        Args:
            directory: 目录路径
            recursive: 是否递归搜索子目录
            tenant_id: 统一注入的租户ID（用于多租户场景）
            department: 统一注入的部门信息
            author: 统一注入的作者信息
            doc_type: 统一注入的文档类型

        Returns:
            Document 列表
        """
        path = Path(directory)
        docs: list[Document] = []

        pattern = "**/*" if recursive else "*"
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self._parsers:
                try:
                    doc = self.load(str(file_path))

                    if tenant_id:
                        doc.metadata["tenant_id"] = tenant_id
                    if department:
                        doc.metadata["department"] = department
                    if author:
                        doc.metadata["author"] = author
                    if doc_type:
                        doc.metadata["doc_type"] = doc_type

                    docs.append(doc)
                except Exception as e:
                    print(f"加载文档失败 {file_path}: {e}")

        return docs
